#!/usr/bin/env python3
"""
Word 报告生成模块 v3.0
生成带删除线、下划线、批注的 .docx 审查报告
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class DocxGenerator:
    """Word 报告生成器"""
    
    def __init__(self):
        self.doc = None
    
    def generate(self, risks: List[Dict], contract_info: Dict, 
                 output_path: str, template_path: Optional[str] = None) -> str:
        """
        生成 Word 格式审查报告
        
        Args:
            risks: 风险项列表
            contract_info: 合同基本信息
            output_path: 输出文件路径
            template_path: 模板文件路径（可选）
            
        Returns:
            生成的文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn, nsmap
            from docx.oxml import OxmlElement
        except ImportError:
            raise ImportError(
                "Word 报告生成需要 python-docx 库。\n"
                "安装方法：pip install python-docx"
            )
        
        # 创建文档
        if template_path and Path(template_path).exists():
            doc = Document(template_path)
        else:
            doc = Document()
        
        self.doc = doc
        
        # 设置默认字体
        self._set_default_font(doc)
        
        # 1. 标题
        title = doc.add_heading('合同智能审查报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. 基本信息表格
        self._add_basic_info(doc, contract_info)
        
        # 3. 风险概况
        self._add_risk_overview(doc, risks, contract_info)
        
        # 4. 风险详情
        self._add_risk_details(doc, risks)
        
        # 5. 缺失条款
        missing = contract_info.get('missing_clauses', [])
        if missing:
            self._add_missing_clauses(doc, missing)
        
        # 6. 修改建议汇总
        if risks:
            self._add_suggestions(doc, risks)
        
        # 7. 免责声明
        self._add_disclaimer(doc)
        
        # 保存
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        logger.info(f"Word 报告已保存: {output_path}")
        return str(output_path)
    
    def _set_default_font(self, doc):
        """设置文档默认字体"""
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)
        # 设置中文字体
        rPr = style.element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            style.element.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _add_basic_info(self, doc, contract_info: Dict):
        """添加基本信息"""
        doc.add_heading('一、基本信息', level=1)
        
        table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(12)
        
        info_items = [
            ('合同名称', contract_info.get('title', '未知')),
            ('合同类型', contract_info.get('contract_type', '未知')),
            ('审查视角', contract_info.get('party_role', '双方')),
            ('审查日期', datetime.now().strftime('%Y-%m-%d')),
        ]
        
        parties = contract_info.get('parties', [])
        if parties:
            party_str = "; ".join([f"{p.get('role', '')}: {p.get('name', '')}" for p in parties])
            if len(party_str) > 100:
                party_str = party_str[:100] + "..."
            info_items.append(('合同主体', party_str))
        
        if contract_info.get('total_amount'):
            info_items.append(('合同金额', f"{contract_info.get('total_amount', '')} {contract_info.get('currency', 'CNY')}"))
        
        for i, (label, value) in enumerate(info_items):
            if i < len(table.rows):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
        
        # 如果行数不够，补充空行
        for i in range(len(info_items), 6):
            row = table.rows[i]
            row.cells[0].text = ''
            row.cells[1].text = ''
        
        doc.add_paragraph('')  # 空行
    
    def _add_risk_overview(self, doc, risks: List[Dict], contract_info: Dict):
        """添加风险概况"""
        doc.add_heading('二、风险概况', level=1)
        
        # 统计
        counts = {'严重': 0, '中等': 0, '一般': 0, '提示': 0}
        for risk in risks:
            severity = risk.get('severity', '一般')
            counts[severity] = counts.get(severity, 0) + 1
        
        # 评分
        score = self._calculate_score(risks)
        score_detail = self._get_score_detail(score)
        
        # 评分和统计表格
        p = doc.add_paragraph()
        run = p.add_run(f'综合评分：{score}/100')
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if score < 50 else \
            RGBColor(0xFF, 0x8C, 0x00) if score < 70 else RGBColor(0x00, 0x80, 0x00)
        
        p = doc.add_paragraph()
        run = p.add_run(f'评分说明：{score_detail}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # 风险等级统计表
        table = doc.add_table(rows=2, cols=4, style='Light Grid Accent 1')
        headers = ['🔴 严重', '🟡 中等', '🟢 一般', 'ℹ️ 提示']
        values = [counts.get('严重', 0), counts.get('中等', 0), counts.get('一般', 0), counts.get('提示', 0)]
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for i, value in enumerate(values):
            cell = table.rows[1].cells[i]
            cell.text = f'{value} 项'
        
        doc.add_paragraph('')
    
    def _add_risk_details(self, doc, risks: List[Dict]):
        """添加风险详情"""
        if not risks:
            doc.add_heading('三、审查结果', level=1)
            doc.add_paragraph('恭喜！本次审查未发现明显风险点。')
            return
        
        doc.add_heading('三、风险详情', level=1)
        
        # 按严重等级排序
        severity_weights = {'严重': 0, '中等': 1, '一般': 2, '提示': 3}
        sorted_risks = sorted(risks, key=lambda r: severity_weights.get(r.get('severity', '一般'), 99))
        
        current_severity = None
        
        for i, risk in enumerate(sorted_risks, 1):
            severity = risk.get('severity', '一般')
            
            if severity != current_severity:
                current_severity = severity
                emoji = {'严重': '🔴', '中等': '🟡', '一般': '🟢', '提示': 'ℹ️'}.get(severity, '⚪')
                doc.add_heading(f'{emoji} {severity}风险', level=2)
            
            # 风险标题（加粗）
            title_para = doc.add_paragraph()
            run = title_para.add_run(f'{i}. {risk.get("title", "未命名风险")}')
            run.bold = True
            run.font.size = Pt(12)
            
            # 位置
            clause_ref = risk.get('clause_ref')
            page = risk.get('page')
            if clause_ref or page:
                location_parts = []
                if clause_ref:
                    location_parts.append(f"条款: {clause_ref}")
                if page:
                    location_parts.append(f"第 {page} 页")
                p = doc.add_paragraph()
                run = p.add_run(f'位置: {", ".join(location_parts)}')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            
            # 原文引用（带删除线 - 表示建议删除/修改）
            if risk.get('text_snippet'):
                p = doc.add_paragraph()
                run = p.add_run('原文: ')
                run.bold = True
                snippet = risk['text_snippet']
                if len(snippet) > 100:
                    snippet = snippet[:97] + "..."
                run = p.add_run(snippet)
                run.strike = True  # 删除线
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            
            # 问题描述
            if risk.get('description'):
                p = doc.add_paragraph()
                run = p.add_run('问题: ')
                run.bold = True
                p.add_run(risk['description'])
            
            # 法律依据
            if risk.get('legal_basis'):
                p = doc.add_paragraph()
                run = p.add_run('法律依据: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
                p.add_run(risk['legal_basis'])
            
            # 修改建议（带下划线 - 表示建议新增/替换的内容）
            if risk.get('suggestion'):
                p = doc.add_paragraph()
                run = p.add_run('修改建议: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                run = p.add_run(risk['suggestion'])
                run.underline = True  # 下划线
                
                # 添加批注
                if len(risk.get('suggestion', '')) > 50:
                    # python-docx 库对批注的支持有限，这里使用文本批注
                    p = doc.add_paragraph()
                    run = p.add_run('📝 批注：建议业务方重点关注此修改点，确保落实到合同条款中。')
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    run.italic = True
            
            # 推荐范本
            if risk.get('template'):
                p = doc.add_paragraph()
                run = p.add_run('推荐范本: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
                p.add_run(risk['template'])
            
            # 分隔线
            p = doc.add_paragraph()
            p.add_run('─' * 50).font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    
    def _add_missing_clauses(self, doc, missing: List[str]):
        """添加缺失条款清单"""
        doc.add_heading('四、缺失条款清单', level=1)
        for item in missing:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'[ ] {item}')
    
    def _add_suggestions(self, doc, risks: List[Dict]):
        """添加修改建议汇总"""
        doc.add_heading('五、修改建议汇总', level=1)
        for i, risk in enumerate(risks, 1):
            if risk.get('suggestion'):
                p = doc.add_paragraph()
                run = p.add_run(f'{i}. {risk.get("title", "")}')
                run.bold = True
                p.add_run(f'\n   → {risk["suggestion"]}')
    
    def _add_disclaimer(self, doc):
        """添加免责声明"""
        doc.add_heading('六、免责声明', level=1)
        
        disclaimers = [
            '本报告由 AI 智能审查工具生成，仅供参考，不构成法律意见。',
            '法律建议需由具备相应执业资格的律师提供。',
            '对于重大合同（标的额超过 100 万元或涉及核心商业利益），强烈建议咨询专业执业律师。',
            '本报告基于当前的法律法规进行审查，法律变更可能导致部分建议过时。',
        ]
        
        for text in disclaimers:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        
        # 反馈邮箱
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run('📬 如有更好建议，欢迎联系：njskills@agent.qq.com')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    
    def _calculate_score(self, risks: List[Dict]) -> int:
        """计算综合评分"""
        score = 100
        score_impact = {'严重': -15, '中等': -8, '一般': -3, '提示': 0}
        for risk in risks:
            severity = risk.get('severity', '一般')
            score += score_impact.get(severity, 0)
        return max(0, min(100, score))
    
    def _get_score_detail(self, score: int) -> str:
        """获取评分说明"""
        score_descriptions = {
            (90, 100): "风险较低，可放心签署",
            (70, 89): "风险中等，建议修改后签署",
            (50, 69): "风险偏高，需要认真修改",
            (0, 49): "风险很高，强烈建议咨询律师",
        }
        for (low, high), desc in sorted(score_descriptions.items(), reverse=True):
            if low <= score <= high:
                return desc
        return "无法评估"


def main():
    """命令行入口"""
    import argparse
    import json
    
    parser = argparse.ArgumentParser(description='Word 报告生成器 v3.0')
    parser.add_argument('risks_json', help='风险列表 JSON 文件路径')
    parser.add_argument('--output', '-o', default='report.docx', help='输出 .docx 路径')
    parser.add_argument('--template', '-t', help='模板文件路径（可选）')
    args = parser.parse_args()
    
    with open(args.risks_json, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    risks = data.get('risks', [])
    contract_info = data.get('contract_info', {})
    
    generator = DocxGenerator()
    output = generator.generate(risks, contract_info, args.output, args.template)
    print(f"✅ Word 报告已生成: {output}")


if __name__ == '__main__':
    main()
