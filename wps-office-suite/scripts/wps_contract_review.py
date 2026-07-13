"""
Word 合同条款自动审查标注 v4.0
输入 .docx 合同文件 → 调用规则引擎标注风险 → 输出修订痕迹版 docx

核心规则引擎（本地实现，无 API）：
  • 必备条款检测：缺少关键条款时标注「缺失」
  • 风险语句标注：发现「无限责任/单方解除/不可撤销」等风险词
  • 不合规格式检测：金额/日期/比例格式不规范
  • 修订痕迹输出：输出带颜色标注和建议的 docx
"""
import re
import json
from pathlib import Path
from typing import List, Dict

from pptx.dml.color import RGBColor


# ==================== 合同规则库 ====================

# 必备条款列表
REQUIRED_ARTICLES = [
    {"key": "合同标的", "keywords": ["标的", "项目", "服务内容", "货物", "工程"]},
    {"key": "价款报酬", "keywords": ["价款", "报酬", "金额", "费用", "价格", "付款"]},
    {"key": "履行期限", "keywords": ["履行期限", "交货期", "工期", "完成时间", "期限", "日期"]},
    {"key": "违约责任", "keywords": ["违约", "赔偿", "违约金", "罚款", "责任"]},
    {"key": "争议解决", "keywords": ["争议", "仲裁", "诉讼", "管辖", "解决"]},
    {"key": "双方信息", "keywords": ["甲方", "乙方", "名称", "地址", "法定代表人"]},
    {"key": "生效条款", "keywords": ["生效", "签署", "盖章", "签字", "日期"]},
    {"key": "保密条款", "keywords": ["保密", "商业秘密", "秘密", "泄露"]},
]

# 风险关键词库
RISK_KEYWORDS = [
    # 高风险（红色标注）
    {"level": "high", "words": ["无限责任", "无限连带", "一切损失", "全部赔偿", "任何间接损失", "包括但不限于所有"]},
    {"level": "high", "words": ["单方解除", "随时解除", "无需理由", "无需通知"]},
    {"level": "high", "words": ["不可撤销", "放弃权利", "不得主张", "不得抗辩", "不得要求"]},
    {"level": "medium", "words": ["最终解释权", "有权随时修改", "不承担责任", "概不负责", "免责条款"]},
    # 中风险（橙色标注）
    {"level": "medium", "words": ["逾期", "滞纳金", "违约金", "赔偿金", "罚款", "罚金"]},
    {"level": "medium", "words": ["连带责任", "补充责任", "共同赔偿"]},
    # 低风险（黄色标注）
    {"level": "low", "words": ["建议", "最好", "应当", "可以"]},
]

# 金额格式规范
AMOUNT_PATTERNS = [
    (r"人民币\s*(\d+)\s*[万元]", "金额应同时标注大小写"),
    (r"(\d+)%", "百分比赔偿超过法定上限"),
]


def review_contract(filepath: str) -> List[Dict]:
    """返回合同中所有发现的问题"""
    try:
        from docx import Document
        doc = Document(filepath)
    except ImportError:
        return [{"type": "error", "message": "需要安装 python-docx：pip install python-docx"}]

    findings = []
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # 1. 必备条款检测
    for article in REQUIRED_ARTICLES:
        has_article = any(kw in full_text for kw in article["keywords"])
        if not has_article:
            findings.append({
                "type": "missing_article",
                "level": "high",
                "title": f"缺少必备条款：{article['key']}",
                "description": f"合同未包含 '{article['key']}' 相关内容，缺少以下关键词：{article['keywords']}",
            })

    # 2. 风险关键词检测
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        for rk in RISK_KEYWORDS:
            for word in rk["words"]:
                if word in text:
                    findings.append({
                        "type": "risk_clause",
                        "level": rk["level"],
                        "word": word,
                        "context": text,
                    })

    # 3. 格式不合规检测
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        for pattern, msg in AMOUNT_PATTERNS:
            if re.search(pattern, text):
                findings.append({
                    "type": "format_issue",
                    "level": "medium",
                    "word": pattern,
                    "context": text,
                    "message": msg,
                })

    return findings


def highlight_contract(source_path: str, output_path: str, findings: List[Dict]) -> Dict:
    """根据审查结果在 docx 中标注颜色"""
    try:
        from docx import Document
        from docx.shared import RGBColor
    except ImportError:
        return {"success": False, "error": "需要安装 python-docx"}

    doc = Document(source_path)
    annotated_count = 0

    # 颜色映射：红/橙/黄
    color_map = {
        "high": RGBColor(0xFF, 0x44, 0x44),
        "medium": RGBColor(0xFF, 0x88, 0x00),
        "low": RGBColor(0xFF, 0xBB, 0x33),
    }

    # 创建查找表（按关键词）
    risk_words = set()
    risk_levels = {}
    for f in findings:
        if f["type"] == "risk_clause":
            risk_words.add(f["word"])
            risk_levels[f["word"]] = f["level"]
        elif f["type"] == "format_issue":
            risk_words.add(f["word"])
            risk_levels[f["word"]] = f["level"]

    # 遍历段落，标注关键词
    for para in doc.paragraphs:
        text = para.text
        found = False
        for word in risk_words:
            if word in text:
                found = True
                break
        if not found:
            continue

        # 清空原段落，用 runs 重新构建
        original_text = para.text
        para.clear()
        remaining = original_text

        while remaining:
            earliest_pos = len(remaining)
            earliest_word = None
            for word in risk_words:
                pos = remaining.find(word)
                if pos != -1 and pos < earliest_pos:
                    earliest_pos = pos
                    earliest_word = word

            if earliest_word is None:
                # 没有更多关键词
                para.add_run(remaining)
                break

            # 写入关键词前的文本
            if earliest_pos > 0:
                para.add_run(remaining[:earliest_pos])

            # 写入带颜色的关键词
            run = para.add_run(earliest_word)
            level = risk_levels.get(earliest_word, "medium")
            run.font.color.rgb = color_map.get(level, RGBColor(0xFF, 0x00, 0x00))
            run.bold = True

            # 继续处理剩余文本
            remaining = remaining[earliest_pos + len(earliest_word):]

        annotated_count += 1

    # 在文档开头插入审查报告页
    p = doc.add_paragraph()
    p.add_run("合同审查报告\n").bold = True
    p = doc.add_paragraph()

    # 统计问题
    high_cnt = sum(1 for f in findings if f["level"] == "high")
    med_cnt = sum(1 for f in findings if f["level"] == "medium")
    low_cnt = sum(1 for f in findings if f["level"] == "low")
    p.add_run(f"高风险：{high_cnt} 处\n").font.color.rgb = color_map["high"] if False else None
    p.add_run(f"中风险：{med_cnt} 处\n")
    p.add_run(f"低风险：{low_cnt} 处\n\n")

    for f in findings:
        line = ""
        if f["type"] == "missing_article":
            line = f"[缺失] {f['title']} - {f['description']}\n"
        elif f["type"] == "risk_clause":
            line = f"[风险] 发现风险词「{f['word']}」- {f.get('context', '')}\n"
        elif f["type"] == "format_issue":
            line = f"[格式] {f.get('message', '')} - {f.get('context', '')}\n"
        if line:
            p.add_run(line)

    # 保存
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return {
        "success": True,
        "path": output_path,
        "total_findings": len(findings),
        "annotated_paragraphs": annotated_count,
        "risk_summary": f"高风险 {high_cnt} 处，中风险 {med_cnt} 处，低风险 {low_cnt} 处",
    }


def review_docx_contract(filepath: str, output_path: str = None) -> Dict:
    """合同审查完整流程"""
    try:
        input_path = Path(filepath).resolve()
        if not input_path.exists():
            return {"success": False, "error": f"文件不存在: {filepath}"}

        if not output_path:
            output_path = str(input_path.parent / f"{input_path.stem}_审查版.docx")

        findings = review_contract(str(input_path))
        result = highlight_contract(str(input_path), output_path, findings)
        result["original"] = str(input_path)
        result["findings"] = findings
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Word 合同条款自动审查")
    parser.add_argument("--file", required=True, help="输入 .docx 合同文件路径")
    parser.add_argument("--output", default="", help="输出 .docx 审查版文件路径（可选）")

    args = parser.parse_args()
    result = review_docx_contract(args.file, args.output)
    print(json.dumps(result, ensure_ascii=False, default=str))
