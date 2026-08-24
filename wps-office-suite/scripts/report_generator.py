"""
报告生成器 v4.9.0
功能：根据要点自动生成结构化工作报告（Word 文档）

v4.9.0 变更：
  - 新增 --polish 参数，通过 llm_bridge 走 cn-llm-router 润色报告内容
  - 未装 cn-llm-router 时 --polish 自动跳过，不影响本地生成

死规则：
  1. 纯本地实现，不依赖任何外部 API
  2. python-docx 未安装时自动降级为 JSON 输出
  3. 所有路径必须经过 safe_path 校验
  4. 中文文档统一使用「微软雅黑」字体
  5. 输出内容以 JSON 格式返回，与其他模块保持一致
  6. 不读取任何外部凭证或配置文件

安全合规：
  - 不联网、不调用外部服务
  - 不读取用户隐私数据
  - 不执行任何系统命令
  - 所有操作仅限于本地文件读写
"""

import sys
import json
import argparse
from pathlib import Path
from datetime import datetime
from typing import Optional, Callable, List, Dict, Any

# 从公共模块导入（带降级）
sys.path.insert(0, str(Path(__file__).parent))

try:
    from wps_common import safe_path
except ImportError:
    def safe_path(path_str: str) -> Path:
        """降级：基本路径校验"""
        path = Path(path_str).resolve()
        parent = path.parent
        if not parent.exists():
            raise FileNotFoundError(f"目录不存在：{parent}")
        return path


# 进度回调默认实现
def _default_progress(step: str, percent: int) -> None:
    """默认进度回调（打印到 stderr）"""
    print(f"[ReportGenerator] {step} ({percent}%)", file=sys.stderr)


class ReportGenerator:
    """工作报告生成器

    支持生成周报、月报等结构化文档，自动排版为正式 Word 格式。
    """

    VERSION = "v4.9.0"
    REPORT_TYPES = {
        "weekly": {
            "label": "周报",
            "title_prefix": "周",
            "sections": {
                "work_summary": "一、本周工作总结",
                "key_progress": "二、重点工作进展",
                "issues": "三、遇到的问题与解决方案",
                "next_plan": "四、下周工作计划",
            },
        },
        "monthly": {
            "label": "月报",
            "title_prefix": "月",
            "sections": {
                "work_summary": "一、本月工作总结",
                "key_progress": "二、重点工作进展",
                "issues": "三、遇到的问题与解决方案",
                "next_plan": "四、下月工作计划",
            },
        },
    }

    TONE_FORMAL = "formal"
    TONE_CASUAL = "casual"

    def __init__(self, progress_callback: Optional[Callable[[str, int], None]] = None):
        """
        Args:
            progress_callback: 进度回调函数，接收 (step, percent) 两个参数
        """
        self.progress = progress_callback or _default_progress

    def generate(
        self,
        report_type: str,
        points: List[str],
        title: str = "",
        author: str = "",
        date: str = "",
        template_path: str = "",
        output: str = "",
        tone: str = "formal",
        polish: bool = False,
    ) -> Dict[str, Any]:
        """生成报告

        Args:
            report_type: 报告类型，"weekly" 或 "monthly"
            points: 关键要点列表
            title: 文档标题（可选）
            author: 作者（可选）
            date: 日期，格式 YYYY-MM-DD（可选，默认今天）
            template_path: 模板文件路径（可选）
            output: 输出 .docx 路径（必填）
            tone: 语气风格，"formal" 或 "casual"（默认 formal）
            polish: 是否通过 llm_bridge 润色内容（默认 False）

        Returns:
            dict: {"ok": bool, "error": str, "path": str, "content": dict}
        """
        try:
            # 校验报告类型
            if report_type not in self.REPORT_TYPES:
                return {
                    "ok": False,
                    "error": f"不支持的报告类型: {report_type}，可选: {list(self.REPORT_TYPES.keys())}",
                }

            # 校验输出路径
            if not output:
                return {"ok": False, "error": "必须指定 --output 输出路径"}

            # 构建内容
            self.progress("构建报告内容", 10)
            if template_path:
                content = self._generate_template_content(template_path, points)
            else:
                content = self._build_content(points, report_type)

            # 填充元数据
            content["title"] = title or f"{self.REPORT_TYPES[report_type]['label']} - {date or datetime.now().strftime('%Y-%m-%d')}"
            content["author"] = author
            content["date"] = date or datetime.now().strftime("%Y-%m-%d")
            content["tone"] = tone
            content["report_type"] = report_type

            # v4.9: 通过 llm_bridge 润色内容（可选）
            if polish:
                content = self._polish_content(content)

            # 生成 Word 文档
            self.progress("生成 Word 文档", 50)
            result = self._generate_docx(content, output)

            if result.get("ok"):
                return {
                    "ok": True,
                    "path": str(output),
                    "content": content,
                    "format": "docx",
                    "version": self.VERSION,
                }
            else:
                # 降级：返回 JSON
                return {
                    "ok": True,
                    "content": content,
                    "format": "json (python-docx 不可用，已降级)",
                    "version": self.VERSION,
                    "warning": result.get("error", ""),
                }

        except Exception as e:
            return {"ok": False, "error": str(e)}

    def _polish_content(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """通过 llm_bridge 润色报告内容（可选）

        Args:
            content: 原始内容结构

        Returns:
            dict: 润色后的内容结构
        """
        try:
            from llm_bridge import summarize as bridge_summarize, is_router_available, INSTALL_HINT
            if not is_router_available():
                content["polish_skipped"] = "cn-llm-router 未安装，跳过润色"
                content["polish_hint"] = INSTALL_HINT
                return content

            # 拼接关键要点进行润色
            points = content.get("sections", {}).get("work_summary", {}).get("items", [])
            if not points:
                return content

            raw_text = "\n".join(f"- {item['text']}" for item in points)
            result = bridge_summarize(raw_text, max_length=200, timeout=30)
            if result.get("ok"):
                content["polished_summary"] = result.get("text", "")
                content["polish_model"] = result.get("model", "")
                content["polish_method"] = "cn-llm-router"
            else:
                content["polish_skipped"] = result.get("error", "润色失败")
        except ImportError:
            content["polish_skipped"] = "llm_bridge 模块不可用"
        except Exception as e:
            content["polish_skipped"] = f"润色异常: {str(e)}"

        return content

    def preview(
        self,
        report_type: str,
        points: List[str],
        title: str = "",
        author: str = "",
        date: str = "",
        tone: str = "formal",
    ) -> Dict[str, Any]:
        """预览报告内容（不生成 docx，直接返回 JSON）

        Args:
            report_type: 报告类型
            points: 关键要点列表
            title: 文档标题（可选）
            author: 作者（可选）
            date: 日期（可选，默认今天）
            tone: 语气风格（默认 formal）

        Returns:
            dict: 结构化内容
        """
        try:
            if report_type not in self.REPORT_TYPES:
                return {
                    "ok": False,
                    "error": f"不支持的报告类型: {report_type}，可选: {list(self.REPORT_TYPES.keys())}",
                }

            content = self._build_content(points, report_type)
            content["title"] = title or f"{self.REPORT_TYPES[report_type]['label']} - {date or datetime.now().strftime('%Y-%m-%d')}"
            content["author"] = author
            content["date"] = date or datetime.now().strftime("%Y-%m-%d")
            content["tone"] = tone
            content["report_type"] = report_type

            return {"ok": True, "content": content, "format": "json", "version": self.VERSION}

        except Exception as e:
            return {"ok": False, "error": str(e)}

    def list_types(self) -> Dict[str, Any]:
        """列出可用的报告类型

        Returns:
            dict: 报告类型信息
        """
        types_info = {}
        for key, info in self.REPORT_TYPES.items():
            types_info[key] = {
                "label": info["label"],
                "sections": list(info["sections"].values()),
            }
        return {"ok": True, "types": types_info, "version": self.VERSION}

    def _build_content(self, points: List[str], report_type: str) -> Dict[str, Any]:
        """构建报告内容结构

        Args:
            points: 关键要点列表
            report_type: 报告类型

        Returns:
            dict: 结构化内容
        """
        type_info = self.REPORT_TYPES[report_type]

        # 工作总结：将每个要点作为编号段落
        summary_items = []
        for i, pt in enumerate(points, 1):
            summary_items.append({"index": i, "text": pt})

        # 重点工作进展表格
        progress_table = []
        for i, pt in enumerate(points, 1):
            progress_table.append({
                "序号": str(i),
                "工作项": pt[:20] + ("..." if len(pt) > 20 else ""),
                "进展": "进行中",
                "状态": "正常",
            })

        content = {
            "sections": {
                "work_summary": {
                    "title": type_info["sections"]["work_summary"],
                    "items": summary_items,
                },
                "key_progress": {
                    "title": type_info["sections"]["key_progress"],
                    "table": {
                        "headers": ["序号", "工作项", "进展", "状态"],
                        "rows": progress_table,
                    },
                },
                "issues": {
                    "title": type_info["sections"]["issues"],
                    "content": "暂无重大问题。",
                },
                "next_plan": {
                    "title": type_info["sections"]["next_plan"],
                    "items": [
                        f"继续推进：{points[0]}" if points else "制定下阶段工作计划",
                    ],
                },
            },
            "attachments": [],
        }

        return content

    def _generate_docx(self, content: Dict[str, Any], output: str) -> Dict[str, Any]:
        """生成 Word 文档

        Args:
            content: 结构化内容
            output: 输出路径

        Returns:
            dict: 生成结果
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            return {
                "ok": False,
                "error": "python-docx 未安装，已降级为 JSON 输出。安装命令: pip install python-docx",
            }

        try:
            # 安全路径校验
            out_path = safe_path(output)

            doc = Document()

            # 设置默认字体（中文：微软雅黑）
            style = doc.styles["Normal"]
            font = style.font
            font.name = "微软雅黑"
            font.size = Pt(11)
            style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 标题（居中、加粗）
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(content.get("title", "工作报告"))
            title_run.bold = True
            title_run.font.size = Pt(18)
            title_run.font.name = "微软雅黑"
            title_run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 日期和作者
            meta_para = doc.add_paragraph()
            meta_text = f"日期：{content.get('date', '')}"
            if content.get("author"):
                meta_text += f"    作者：{content['author']}"
            meta_run = meta_para.add_run(meta_text)
            meta_run.font.size = Pt(10)
            meta_run.font.name = "微软雅黑"
            meta_run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            doc.add_paragraph()  # 空行

            # 一、工作总结
            sections = content.get("sections", {})
            summary = sections.get("work_summary", {})
            doc.add_heading(summary.get("title", "一、工作总结"), level=1)
            for item in summary.get("items", []):
                p = doc.add_paragraph(f"{item['index']}. {item['text']}")
                p.style = doc.styles["List Number"]
                for run in p.runs:
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 二、重点工作进展（表格）
            progress = sections.get("key_progress", {})
            doc.add_heading(progress.get("title", "二、重点工作进展"), level=1)
            table_data = progress.get("table", {})
            headers = table_data.get("headers", ["序号", "工作项", "进展", "状态"])
            rows = table_data.get("rows", [])

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"

            # 表头
            for j, header in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.name = "微软雅黑"
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 表体
            for i, row_data in enumerate(rows):
                for j, header in enumerate(headers):
                    cell = table.rows[i + 1].cells[j]
                    cell.text = row_data.get(header, "")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = "微软雅黑"
                            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            doc.add_paragraph()  # 空行

            # 三、遇到的问题与解决方案
            issues = sections.get("issues", {})
            doc.add_heading(issues.get("title", "三、遇到的问题与解决方案"), level=1)
            p = doc.add_paragraph(issues.get("content", "暂无。"))
            for run in p.runs:
                run.font.name = "微软雅黑"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 四、下阶段工作计划
            next_plan = sections.get("next_plan", {})
            doc.add_heading(next_plan.get("title", "四、下阶段工作计划"), level=1)
            for item in next_plan.get("items", []):
                p = doc.add_paragraph(item)
                p.style = doc.styles["List Bullet"]
                for run in p.runs:
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 附件说明（可选）
            attachments = content.get("attachments", [])
            if attachments:
                doc.add_heading("附件说明", level=1)
                for att in attachments:
                    doc.add_paragraph(att, style=doc.styles["List Bullet"])

            # 保存文档
            doc.save(str(out_path))

            return {"ok": True, "path": str(out_path)}

        except Exception as e:
            return {"ok": False, "error": f"DOCX 生成失败: {str(e)}"}

    def _generate_template_content(self, template_path: str, points: List[str]) -> Dict[str, Any]:
        """基于模板文件生成内容

        Args:
            template_path: 模板文件路径（.docx 或 .json）
            points: 关键要点列表

        Returns:
            dict: 结构化内容
        """
        try:
            t_path = safe_path(template_path)
            suffix = t_path.suffix.lower()

            if suffix == ".json":
                template = json.loads(t_path.read_text(encoding="utf-8"))
                report_type = template.get("report_type", "weekly")
                type_info = self.REPORT_TYPES.get(report_type, self.REPORT_TYPES["weekly"])

                content = self._build_content(points, report_type)
                content["title"] = template.get("title", content.get("title", ""))
                content["author"] = template.get("author", "")
                content["tone"] = template.get("tone", "formal")
                content["report_type"] = report_type

                # 合并模板中的自定义内容
                for key in ["issues", "next_plan"]:
                    if key in template.get("sections", {}):
                        content["sections"][key].update(template["sections"][key])

                if template.get("attachments"):
                    content["attachments"] = template["attachments"]

                return content

            elif suffix == ".docx":
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(str(t_path))
                    report_type = "weekly"
                    type_info = self.REPORT_TYPES[report_type]

                    # 提取模板中的纯文本要点
                    template_points = []
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text and not text.startswith("一、") and not text.startswith("二、") and not text.startswith("三、") and not text.startswith("四、"):
                            if text not in [type_info["sections"]["work_summary"], type_info["sections"]["key_progress"], type_info["sections"]["issues"], type_info["sections"]["next_plan"]]:
                                template_points.append(text)

                    # 合并模板要点与传入要点
                    all_points = template_points + points
                    content = self._build_content(all_points, report_type)

                    # 从模板提取标题
                    for para in doc.paragraphs:
                        if para.text.strip() and para.runs and para.runs[0].bold:
                            content["title"] = para.text.strip()
                            break

                    return content

                except ImportError:
                    # python-docx 不可用时，降级
                    return self._build_content(points, "weekly")

            else:
                return self._build_content(points, "weekly")

        except Exception as e:
            # 模板处理失败，降级为默认构建
            print(f"[ReportGenerator] 模板处理失败，使用默认构建: {e}", file=sys.stderr)
            return self._build_content(points, "weekly")


def main():
    """CLI 入口"""
    parser = argparse.ArgumentParser(
        description=f"工作报告生成器 {ReportGenerator.VERSION}",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)

    # generate 子命令
    p_gen = sub.add_parser("generate", help="生成工作报告（Word 文档）")
    p_gen.add_argument("--type", required=True, choices=["weekly", "monthly"], help="报告类型")
    p_gen.add_argument("--points", required=True, help="关键要点，逗号分隔")
    p_gen.add_argument("--title", default="", help="文档标题")
    p_gen.add_argument("--author", default="", help="作者")
    p_gen.add_argument("--date", default="", help="日期 YYYY-MM-DD（默认今天）")
    p_gen.add_argument("--template", default="", help="模板文件路径（.docx 或 .json）")
    p_gen.add_argument("--output", required=True, help="输出 .docx 路径")
    p_gen.add_argument("--tone", default="formal", choices=["formal", "casual"], help="语气风格")
    p_gen.add_argument("--polish", action="store_true", help="通过 llm_bridge 润色内容（需 cn-llm-router）")

    # check 子命令
    p_check = sub.add_parser("check", help="查看可用的报告类型")

    # preview 子命令
    p_prev = sub.add_parser("preview", help="预览报告内容（JSON，不生成 docx）")
    p_prev.add_argument("--type", required=True, choices=["weekly", "monthly"], help="报告类型")
    p_prev.add_argument("--points", required=True, help="关键要点，逗号分隔")
    p_prev.add_argument("--title", default="", help="文档标题")
    p_prev.add_argument("--author", default="", help="作者")
    p_prev.add_argument("--date", default="", help="日期 YYYY-MM-DD（默认今天）")
    p_prev.add_argument("--tone", default="formal", choices=["formal", "casual"], help="语气风格")

    args = parser.parse_args()

    gen = ReportGenerator()

    if args.command == "generate":
        points = [p.strip() for p in args.points.split(",") if p.strip()]
        if not points:
            r = {"ok": False, "error": "要点不能为空"}
        else:
            r = gen.generate(
                report_type=args.type,
                points=points,
                title=args.title,
                author=args.author,
                date=args.date,
                template_path=args.template,
                output=args.output,
                tone=args.tone,
                polish=args.polish,
            )
    elif args.command == "check":
        r = gen.list_types()
    elif args.command == "preview":
        points = [p.strip() for p in args.points.split(",") if p.strip()]
        if not points:
            r = {"ok": False, "error": "要点不能为空"}
        else:
            r = gen.preview(
                report_type=args.type,
                points=points,
                title=args.title,
                author=args.author,
                date=args.date,
                tone=args.tone,
            )
    else:
        r = {"ok": False, "error": "未知命令"}

    print(json.dumps(r, ensure_ascii=False, default=str))


if __name__ == "__main__":
    main()
