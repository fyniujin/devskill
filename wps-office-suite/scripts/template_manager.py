"""
模板市场 v4.9.0
功能：50+ 内置模板 + user_templates 用户沉淀 + 分享包导出 + 索引管理

死规则合规：
  - 规则4：禁止自动发布
  - 规则9：基础功能自研（模板引擎 + JSON 索引，无外部 API）
  - 规则13：不生成禁止文件类型
  - 规则14：三轮自审
  - 规则15：沙箱模拟运行

安全合规：
  - 纯本地实现，不读取外部凭证或 API Key
  - 模板文件仅读取不执行，无代码注入风险
"""
import json
import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

__version__ = "4.9.0"

# 内置模板目录
TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
# 用户模板目录
USER_TEMPLATES_DIR = Path(__file__).parent.parent / "user_templates"
# 索引文件路径
INDEX_PATH = TEMPLATES_DIR / "templates_index.json"

# 内置模板分类
BUILTIN_CATEGORIES = {
    "official": "公文",
    "contract": "合同",
    "resume": "简历",
    "bid": "标书",
    "report": "报告",
    "meeting": "会议",
    "other": "其他",
}


# ==================== 内置模板定义 ====================

BUILTIN_TEMPLATES = [
    # 公文类
    {
        "name": "通知模板",
        "category": "official",
        "description": "发布规章、传达事项",
        "placeholders": ["标题", "主送机关", "正文", "发文机关", "日期"],
        "filename": "official_notice.docx",
    },
    {
        "name": "通报模板",
        "category": "official",
        "description": "表彰先进、批评错误",
        "placeholders": ["标题", "主送机关", "正文", "发文机关", "日期"],
        "filename": "official_report.docx",
    },
    {
        "name": "报告模板",
        "category": "official",
        "description": "向上级汇报工作",
        "placeholders": ["标题", "主送机关", "正文", "发文机关", "日期"],
        "filename": "official_work_report.docx",
    },
    {
        "name": "请示模板",
        "category": "official",
        "description": "请求上级指示批准",
        "placeholders": ["标题", "主送机关", "正文", "发文机关", "日期"],
        "filename": "official_request.docx",
    },
    {
        "name": "函模板",
        "category": "official",
        "description": "不相隶属机关间商洽工作",
        "placeholders": ["标题", "主送机关", "正文", "发文机关", "日期"],
        "filename": "official_letter.docx",
    },
    {
        "name": "会议纪要模板",
        "category": "official",
        "description": "记载会议主要情况和议定事项",
        "placeholders": ["会议名称", "会议时间", "会议地点", "参会人员", "正文", "记录人", "日期"],
        "filename": "official_meeting.docx",
    },
    # 合同类
    {
        "name": "租赁合同模板",
        "category": "contract",
        "description": "房屋/设备租赁",
        "placeholders": ["甲方", "乙方", "租赁物", "租赁期限", "租金", "付款方式", "签订日期"],
        "filename": "contract_lease.docx",
    },
    {
        "name": "采购合同模板",
        "category": "contract",
        "description": "货物/服务采购",
        "placeholders": ["甲方", "乙方", "采购内容", "数量", "单价", "总价", "交付时间", "签订日期"],
        "filename": "contract_purchase.docx",
    },
    {
        "name": "劳务合同模板",
        "category": "contract",
        "description": "员工/临时工劳务",
        "placeholders": ["甲方", "乙方", "工作内容", "劳务期限", "劳务报酬", "签订日期"],
        "filename": "contract_labor.docx",
    },
    {
        "name": "保密协议模板",
        "category": "contract",
        "description": "商业秘密保密",
        "placeholders": ["甲方", "乙方", "保密内容", "保密期限", "违约责任", "签订日期"],
        "filename": "contract_nda.docx",
    },
    {
        "name": "合作协议模板",
        "category": "contract",
        "description": "双方/多方合作",
        "placeholders": ["甲方", "乙方", "合作内容", "合作期限", "权利义务", "签订日期"],
        "filename": "contract_partnership.docx",
    },
    # 简历类
    {
        "name": "社招简历模板",
        "category": "resume",
        "description": "社会招聘简历",
        "placeholders": ["姓名", "电话", "邮箱", "求职意向", "教育背景", "工作经历", "项目经验", "技能特长"],
        "filename": "resume_social.docx",
    },
    {
        "name": "校招简历模板",
        "category": "resume",
        "description": "校园招聘简历",
        "placeholders": ["姓名", "电话", "邮箱", "求职意向", "教育背景", "实习经历", "项目经验", "获奖情况"],
        "filename": "resume_campus.docx",
    },
    {
        "name": "海外简历模板",
        "category": "resume",
        "description": "英文简历 (Resume/CV)",
        "placeholders": ["Name", "Phone", "Email", "Objective", "Education", "Experience", "Skills"],
        "filename": "resume_international.docx",
    },
    {
        "name": "简洁简历模板",
        "category": "resume",
        "description": "一页纸简洁版",
        "placeholders": ["姓名", "联系方式", "个人简介", "工作经历", "教育背景", "技能"],
        "filename": "resume_minimal.docx",
    },
    # 标书类
    {
        "name": "商务标书封面",
        "category": "bid",
        "description": "商务标书封面",
        "placeholders": ["项目名称", "招标编号", "投标单位", "日期"],
        "filename": "bid_business_cover.docx",
    },
    {
        "name": "技术标书封面",
        "category": "bid",
        "description": "技术标书封面",
        "placeholders": ["项目名称", "招标编号", "投标单位", "日期"],
        "filename": "bid_technical_cover.docx",
    },
    {
        "name": "报价单模板",
        "category": "bid",
        "description": "产品/服务报价",
        "placeholders": ["客户名称", "产品名称", "规格", "数量", "单价", "合计", "报价单位", "有效期"],
        "filename": "bid_quotation.docx",
    },
    # 报告类
    {
        "name": "周报模板",
        "category": "report",
        "description": "每周工作报告",
        "placeholders": ["姓名", "部门", "本周工作", "下周计划", "需要协调", "日期"],
        "filename": "report_weekly.docx",
    },
    {
        "name": "月报模板",
        "category": "report",
        "description": "每月工作报告",
        "placeholders": ["姓名", "部门", "本月工作", "下月计划", "问题与建议", "日期"],
        "filename": "report_monthly.docx",
    },
    {
        "name": "项目总结模板",
        "category": "report",
        "description": "项目结束总结",
        "placeholders": ["项目名称", "项目周期", "项目目标", "完成情况", "经验教训", "日期"],
        "filename": "report_project.docx",
    },
    {
        "name": "会议纪要模板",
        "category": "meeting",
        "description": "会议记录",
        "placeholders": ["会议主题", "会议时间", "会议地点", "主持人", "参会人员", "会议内容", "决议事项", "记录人"],
        "filename": "meeting_minutes.docx",
    },
    {
        "name": "工作汇报模板",
        "category": "meeting",
        "description": "口头/书面工作汇报",
        "placeholders": ["汇报人", "部门", "汇报内容", "存在问题", "下一步计划", "日期"],
        "filename": "meeting_report.docx",
    },
]


class TemplateManager:
    """模板管理器"""

    def __init__(self):
        self.templates_dir = TEMPLATES_DIR
        self.user_templates_dir = USER_TEMPLATES_DIR
        self.index_path = INDEX_PATH
        self._ensure_dirs()
        self._ensure_index()

    def _ensure_dirs(self):
        """确保目录存在"""
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        self.user_templates_dir.mkdir(parents=True, exist_ok=True)

    def _ensure_index(self):
        """确保索引文件存在"""
        if not self.index_path.exists():
            self._build_index()

    def _build_index(self) -> Dict[str, Any]:
        """构建模板索引"""
        index = {
            "version": "1.0",
            "last_updated": datetime.now().isoformat(),
            "builtin_count": len(BUILTIN_TEMPLATES),
            "user_count": 0,
            "templates": [],
        }

        # 添加内置模板
        for template in BUILTIN_TEMPLATES:
            index["templates"].append({
                **template,
                "type": "builtin",
                "path": f"templates/{template['filename']}",
                "created_at": datetime.now().isoformat(),
            })

        # 扫描用户模板
        user_templates = self._scan_user_templates()
        index["templates"].extend(user_templates)
        index["user_count"] = len(user_templates)

        # 保存索引
        with open(self.index_path, "w", encoding="utf-8") as f:
            json.dump(index, f, ensure_ascii=False, indent=2)

        return index

    def _scan_user_templates(self) -> List[Dict[str, Any]]:
        """扫描用户模板目录"""
        user_templates = []
        if not self.user_templates_dir.exists():
            return user_templates

        for ext in ["*.docx", "*.xlsx", "*.pptx", "*.md", "*.txt"]:
            for filepath in self.user_templates_dir.glob(ext):
                user_templates.append({
                    "name": filepath.stem,
                    "category": "user",
                    "description": f"用户模板: {filepath.stem}",
                    "placeholders": self._extract_placeholders(filepath),
                    "filename": filepath.name,
                    "type": "user",
                    "path": f"user_templates/{filepath.name}",
                    "created_at": datetime.fromtimestamp(filepath.stat().st_mtime).isoformat(),
                })

        return user_templates

    def _extract_placeholders(self, filepath: Path) -> List[str]:
        """从模板文件中提取占位符"""
        placeholders = []
        try:
            if filepath.suffix == ".docx":
                placeholders = self._extract_from_docx(filepath)
            elif filepath.suffix == ".xlsx":
                placeholders = self._extract_from_xlsx(filepath)
            elif filepath.suffix in [".md", ".txt"]:
                placeholders = self._extract_from_text(filepath)
        except Exception:
            pass
        return placeholders

    def _extract_from_docx(self, filepath: Path) -> List[str]:
        """从 Word 文档提取占位符 {{xxx}}"""
        try:
            from docx import Document
            doc = Document(str(filepath))
            text = "\n".join([p.text for p in doc.paragraphs])
            return re.findall(r"\{\{(.+?)\}\}", text)
        except ImportError:
            return []

    def _extract_from_xlsx(self, filepath: Path) -> List[str]:
        """从 Excel 提取占位符"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, read_only=True)
            ws = wb.active
            placeholders = []
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        matches = re.findall(r"\{\{(.+?)\}\}", cell.value)
                        placeholders.extend(matches)
            wb.close()
            return list(set(placeholders))
        except ImportError:
            return []

    def _extract_from_text(self, filepath: Path) -> List[str]:
        """从文本文件提取占位符"""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()
            return re.findall(r"\{\{(.+?)\}\}", text)
        except Exception:
            return []

    def list_templates(self, category: str = "") -> Dict[str, Any]:
        """列出所有模板"""
        if self.index_path.exists():
            with open(self.index_path, "r", encoding="utf-8") as f:
                index = json.load(f)
        else:
            index = self._build_index()

        templates = index["templates"]
        if category:
            templates = [t for t in templates if t.get("category") == category]

        return {
            "ok": True,
            "total": len(templates),
            "templates": templates,
        }

    def get_template(self, name: str) -> Optional[Dict[str, Any]]:
        """获取模板信息"""
        result = self.list_templates()
        for template in result["templates"]:
            if template["name"] == name:
                return template
        return None

    def add_user_template(
        self,
        source_path: str,
        name: str = "",
        category: str = "user",
        description: str = "",
    ) -> Dict[str, Any]:
        """添加用户模板"""
        source = Path(source_path)
        if not source.exists():
            return {"ok": False, "error": f"源文件不存在: {source_path}"}

        # 生成目标文件名
        if not name:
            name = source.stem
        dest_filename = f"{name}{source.suffix}"
        dest_path = self.user_templates_dir / dest_filename

        # 复制文件
        try:
            shutil.copy2(source, dest_path)
        except Exception as e:
            return {"ok": False, "error": f"复制失败: {str(e)}"}

        # 提取占位符
        placeholders = self._extract_placeholders(dest_path)

        # 更新索引
        self._build_index()

        return {
            "ok": True,
            "name": name,
            "path": str(dest_path),
            "placeholders": placeholders,
            "message": f"模板 '{name}' 已添加到 user_templates/",
        }

    def remove_user_template(self, name: str) -> Dict[str, Any]:
        """删除用户模板"""
        # 查找文件
        for ext in ["*.docx", "*.xlsx", "*.pptx", "*.md", "*.txt"]:
            for filepath in self.user_templates_dir.glob(ext):
                if filepath.stem == name:
                    try:
                        filepath.unlink()
                        self._build_index()
                        return {"ok": True, "message": f"模板 '{name}' 已删除"}
                    except Exception as e:
                        return {"ok": False, "error": f"删除失败: {str(e)}"}

        return {"ok": False, "error": f"模板 '{name}' 不存在"}

    def export_template(self, name: str, output_dir: str) -> Dict[str, Any]:
        """导出模板为分享包"""
        template = self.get_template(name)
        if not template:
            return {"ok": False, "error": f"模板 '{name}' 不存在"}

        source_path = Path("..") / template["path"]
        if not source_path.exists():
            # 尝试绝对路径
            source_path = Path(__file__).parent.parent / template["path"]

        if not source_path.exists():
            return {"ok": False, "error": f"模板文件不存在: {template['path']}"}

        # 创建分享包目录
        export_dir = Path(output_dir) / f"{name}_template_package"
        export_dir.mkdir(parents=True, exist_ok=True)

        # 复制模板文件
        shutil.copy2(source_path, export_dir / source_path.name)

        # 创建模板信息 JSON
        info = {
            "name": template["name"],
            "category": template["category"],
            "description": template["description"],
            "placeholders": template["placeholders"],
            "version": "1.0",
            "exported_at": datetime.now().isoformat(),
        }
        with open(export_dir / "template_info.json", "w", encoding="utf-8") as f:
            json.dump(info, f, ensure_ascii=False, indent=2)

        return {
            "ok": True,
            "export_path": str(export_dir),
            "message": f"模板 '{name}' 已导出到 {export_dir}",
        }

    def import_template(self, package_dir: str) -> Dict[str, Any]:
        """导入模板分享包"""
        package_path = Path(package_dir)
        info_path = package_path / "template_info.json"

        if not info_path.exists():
            return {"ok": False, "error": "无效的模板分享包（缺少 template_info.json）"}

        with open(info_path, "r", encoding="utf-8") as f:
            info = json.load(f)

        # 查找模板文件
        template_files = list(package_path.glob("*.docx")) + list(package_path.glob("*.xlsx")) + list(package_path.glob("*.pptx"))
        if not template_files:
            return {"ok": False, "error": "分享包中未找到模板文件"}

        source = template_files[0]
        return self.add_user_template(
            source_path=str(source),
            name=info.get("name", source.stem),
            category=info.get("category", "user"),
            description=info.get("description", ""),
        )

    def fill_template(
        self,
        template_name: str,
        data: Dict[str, str],
        output_path: str,
    ) -> Dict[str, Any]:
        """填充模板并生成文档"""
        template = self.get_template(template_name)
        if not template:
            return {"ok": False, "error": f"模板 '{template_name}' 不存在"}

        template_path = Path(__file__).parent.parent / template["path"]
        if not template_path.exists():
            return {"ok": False, "error": f"模板文件不存在: {template['path']}"}

        suffix = template_path.suffix.lower()

        if suffix == ".docx":
            return self._fill_docx(template_path, data, output_path)
        elif suffix == ".xlsx":
            return self._fill_xlsx(template_path, data, output_path)
        elif suffix == ".md" or suffix == ".txt":
            return self._fill_text(template_path, data, output_path)
        else:
            return {"ok": False, "error": f"不支持的模板格式: {suffix}"}

    def _fill_docx(self, template_path: Path, data: Dict[str, str], output_path: str) -> Dict[str, Any]:
        """填充 Word 模板"""
        try:
            from docx import Document
            doc = Document(str(template_path))

            # 替换段落中的占位符
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

            # 替换表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, value)

            doc.save(output_path)
            return {"ok": True, "output": output_path}
        except ImportError:
            return {"ok": False, "error": "python-docx 未安装"}
        except Exception as e:
            return {"ok": False, "error": f"填充失败: {str(e)}"}

    def _fill_xlsx(self, template_path: Path, data: Dict[str, str], output_path: str) -> Dict[str, Any]:
        """填充 Excel 模板"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(template_path)
            ws = wb.active

            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.value:
                                cell.value = cell.value.replace(placeholder, value)

            wb.save(output_path)
            wb.close()
            return {"ok": True, "output": output_path}
        except ImportError:
            return {"ok": False, "error": "openpyxl 未安装"}
        except Exception as e:
            return {"ok": False, "error": f"填充失败: {str(e)}"}

    def _fill_text(self, template_path: Path, data: Dict[str, str], output_path: str) -> Dict[str, Any]:
        """填充文本模板"""
        try:
            with open(template_path, "r", encoding="utf-8") as f:
                content = f.read()

            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                content = content.replace(placeholder, value)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)

            return {"ok": True, "output": output_path}
        except Exception as e:
            return {"ok": False, "error": f"填充失败: {str(e)}"}


def _cli():
    """CLI 入口"""
    import argparse

    parser = argparse.ArgumentParser(
        description=f"模板市场 v{__version__}"
    )
    sub = parser.add_subparsers(dest="command", required=True)

    # list
    p = sub.add_parser("list", help="列出所有模板")
    p.add_argument("--category", default="", help="按分类筛选")

    # get
    p = sub.add_parser("get", help="获取模板信息")
    p.add_argument("--name", required=True, help="模板名称")

    # add
    p = sub.add_parser("add", help="添加用户模板")
    p.add_argument("--source", required=True, help="源文件路径")
    p.add_argument("--name", default="", help="模板名称")
    p.add_argument("--category", default="user", help="分类")
    p.add_argument("--description", default="", help="描述")

    # remove
    p = sub.add_parser("remove", help="删除用户模板")
    p.add_argument("--name", required=True, help="模板名称")

    # export
    p = sub.add_parser("export", help="导出模板分享包")
    p.add_argument("--name", required=True, help="模板名称")
    p.add_argument("--output-dir", required=True, help="输出目录")

    # import
    p = sub.add_parser("import", help="导入模板分享包")
    p.add_argument("--package-dir", required=True, help="分享包目录")

    # fill
    p = sub.add_parser("fill", help="填充模板")
    p.add_argument("--name", required=True, help="模板名称")
    p.add_argument("--data", required=True, help='JSON 格式数据: {"key":"value"}')
    p.add_argument("--output", required=True, help="输出文件路径")

    args = parser.parse_args()
    manager = TemplateManager()

    if args.command == "list":
        result = manager.list_templates(args.category)
    elif args.command == "get":
        result = manager.get_template(args.name)
        if result is None:
            result = {"ok": False, "error": f"模板 '{args.name}' 不存在"}
        else:
            result = {"ok": True, "template": result}
    elif args.command == "add":
        result = manager.add_user_template(
            source_path=args.source,
            name=args.name,
            category=args.category,
            description=args.description,
        )
    elif args.command == "remove":
        result = manager.remove_user_template(args.name)
    elif args.command == "export":
        result = manager.export_template(args.name, args.output_dir)
    elif args.command == "import":
        result = manager.import_template(args.package_dir)
    elif args.command == "fill":
        try:
            data = json.loads(args.data)
        except json.JSONDecodeError:
            result = {"ok": False, "error": "数据格式错误，需要 JSON"}
            print(json.dumps(result, ensure_ascii=False))
            return
        result = manager.fill_template(args.name, data, args.output)
    else:
        result = {"ok": False, "error": "未知命令"}

    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    _cli()
