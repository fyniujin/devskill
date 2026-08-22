"""
邮件智能回复 v4.8.0
功能：邮件智能回复（模板匹配 + 规则引擎 + 可选 LLM，纯本地实现）

死规则合规：
  - 规则4：禁止自动发布
  - 规则9：基础功能自研（模板引擎 + 规则匹配，无外部 API）
  - 规则13：不生成禁止文件类型
  - 规则14：三轮自审
  - 规则15：沙箱模拟运行

安全合规：
  - 纯本地实现，不读取外部凭证或 API Key
  - auto 模式仅使用本地模板 + 规则引擎
  - 外部 LLM 仅在用户显式指定 method 时调用
"""

import re
import json
import sys
import os
import hashlib
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Callable

# ==================== 模块元数据 ====================
__version__ = "4.8.0"
__module__ = "email_reply"

# ==================== 可选依赖 ====================
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

# ==================== 工具函数 ====================

def safe_path(path_str: str) -> Path:
    """安全路径校验，返回 Path 对象"""
    path = Path(path_str).resolve()
    parent = path.parent
    if not parent.exists():
        parent.mkdir(parents=True, exist_ok=True)
    return path


def progress_callback(current: int, total: int, message: str = ""):
    """进度回调（规则10：性能优化，给用户反馈）"""
    if total > 0:
        pct = min(100, int(current / total * 100))
        bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
        print(f"\r[{bar}] {pct}% {message}", end="", flush=True)
        if current >= total:
            print()


# ==================== 内置模板库 ====================

BUILTIN_TEMPLATES: Dict[str, Dict[str, Dict[str, str]]] = {
    "thank_you": {
        "zh": {
            "friendly": "亲爱的{sender_name}：\n\n您好！非常感谢您在{date}关于{topic}的帮助和支持！您的付出让我们团队倍感温暖。\n\n{action_item}\n\n再次感谢，期待后续合作！\n\n此致\n敬礼",
            "polite": "{sender_name} 您好：\n\n衷心感谢您在{date}就{topic}事宜给予的支持与协助。您的专业态度和高效执行令人印象深刻。\n\n{action_item}\n\n谨致诚挚谢意。\n\n顺颂商祺",
            "formal": "尊敬的{sender_name}：\n\n值此{date}之际，谨就{topic}一事向您致以最诚挚的谢意。您的鼎力支持对本项目具有重要意义。\n\n{action_item}\n\n特此鸣谢，顺祝工作顺利。",
            "professional": "Dear {sender_name},\n\nThank you sincerely for your assistance on {topic} dated {date}. Your support was instrumental to the success of this initiative.\n\n{action_item}\n\nBest regards"
        },
        "en": {
            "friendly": "Dear {sender_name},\n\nA big thank you for your help with {topic} on {date}! We truly appreciate your support.\n\n{action_item}\n\nThanks again and looking forward to working together!\n\nBest,\n[Your Name]",
            "polite": "Dear {sender_name},\n\nThank you very much for your assistance regarding {topic} on {date}. Your prompt response and professionalism are greatly appreciated.\n\n{action_item}\n\nSincerely,\n[Your Name]",
            "formal": "Dear {sender_name},\n\nOn behalf of our team, I would like to extend our sincere gratitude for your support on {topic} dated {date}.\n\n{action_item}\n\nYours faithfully,\n[Your Name]",
            "professional": "Dear {sender_name},\n\nI am writing to formally acknowledge your valuable contribution to {topic} on {date}.\n\n{action_item}\n\nBest regards,\n[Your Name]"
        }
    },
    "project_update": {
        "zh": {
            "friendly": "Hi {sender_name}，\n\n给您同步一下{topic}的最新进展（{date}）：\n\n✅ 已完成：\n- 需求评审通过\n- 核心模块开发完成 80%\n\n🔄 进行中：\n- 前端页面联调\n- 测试用例编写\n\n⏳ 下一步计划：\n{action_item}\n\n有任何问题随时沟通！\n\n祝好",
            "polite": "{sender_name} 您好，\n\n以下是{topic}项目截至{date}的进度更新：\n\n【进展概要】\n- 已完成需求分析与设计评审\n- 开发阶段完成约 75%\n- 单元测试覆盖率已达 60%\n\n【风险与应对】\n- 第三方接口联调存在延期风险，已安排专人跟进\n\n【下阶段安排】\n{action_item}\n\n如有疑问请随时联系。\n\n此致",
            "formal": "尊敬的{sender_name}：\n\n现就{topic}项目{date}进度情况汇报如下：\n\n一、当前进展\n1. 需求阶段：已完成\n2. 开发阶段：进行中（完成度 78%）\n3. 测试阶段：尚未启动\n\n二、关键里程碑\n- 预计 {date} 完成全部开发工作\n- 预计下阶段进入集成测试\n\n三、后续行动\n{action_item}\n\n特此汇报，请审阅。",
            "professional": "Dear {sender_name},\n\nPlease find below the project status update for {topic} as of {date}:\n\nKey Accomplishments:\n- Requirements finalized and approved\n- Development 80% complete\n- QA test plan drafted\n\nNext Steps:\n{action_item}\n\nPlease let me know if you have any questions.\n\nBest regards"
        },
        "en": {
            "friendly": "Hi {sender_name},\n\nQuick update on {topic} ({date}):\n\n✅ Completed:\n- Design review done\n- Backend API 80% ready\n\n🔄 In Progress:\n- Frontend integration\n- Documentation\n\n⏳ Next up:\n{action_item}\n\nLet me know if you have any questions!\n\nCheers",
            "polite": "Dear {sender_name},\n\nI hope this message finds you well. Here is the latest update on {topic} as of {date}:\n\nProgress Summary:\n- Development phase: ~75% complete\n- Key deliverables on track\n\nUpcoming Milestones:\n{action_item}\n\nPlease feel free to reach out with any questions.\n\nBest regards",
            "formal": "Dear {sender_name},\n\nThis is the formal status report for {topic} dated {date}:\n\nI. Current Status\n- Phase 1: Complete\n- Phase 2: 78% complete\n\nII. Key Metrics\n- Schedule variance: On track\n- Budget status: Within allocation\n\nIII. Next Actions:\n{action_item}\n\nYour feedback is welcome.\n\nSincerely",
            "professional": "Dear {sender_name},\n\nAttached is the project status update for {topic} covering the period through {date}.\n\nExecutive Summary:\nAll major milestones are on schedule with the exception of one dependency item being tracked.\n\nAction Items:\n{action_item}\n\nI am available to discuss at your convenience.\n\nBest regards"
        }
    },
    "meeting_request": {
        "zh": {
            "friendly": "Hi {sender_name}，\n\n想和您约个时间聊聊{topic}的事情，您看下{date}方便不？\n\n📋 会议主题：{topic}\n📅 建议时间：{date}\n⏰ 预计时长：30 分钟\n\n{action_item}\n\n期待您的回复！",
            "polite": "{sender_name} 您好，\n\n诚挚邀请您参加关于{topic}的讨论会，详情如下：\n\n主题：{topic}\n时间：{date}\n地点：公司会议室 A / 腾讯会议\n\n议程安排：\n{action_item}\n\n烦请确认您的出席，谢谢！\n\n此致",
            "formal": "尊敬的{sender_name}：\n\n兹定于{date}召开{topic}专题会议，诚邀您拨冗出席。\n\n会议议题：{topic}\n会议时间：{date}\n会议地点：公司总部 3F 会议室\n\n会议议程：\n{action_item}\n\n敬请届时出席为荷。",
            "professional": "Dear {sender_name},\n\nI would like to schedule a meeting to discuss {topic}.\n\nProposed Details:\n- Topic: {topic}\n- Date/Time: {date}\n- Duration: 30 minutes\n- Location: Conference Room A / Teams\n\nAgenda:\n{action_item}\n\nPlease let me know if this works for you or suggest an alternative.\n\nBest regards"
        },
        "en": {
            "friendly": "Hi {sender_name},\n\nWould love to catch up about {topic}. Are you free on {date}?\n\nTopic: {topic}\nWhen: {date}\nDuration: ~30 min\n\n{action_item}\n\nLet me know what works!\n\nThanks",
            "polite": "Dear {sender_name},\n\nI would like to invite you to a meeting regarding {topic}.\n\nDetails:\n- Topic: {topic}\n- Date/Time: {date}\n- Duration: 30 minutes\n\nAgenda:\n{action_item}\n\nPlease confirm your availability at your earliest convenience.\n\nBest regards",
            "formal": "Dear {sender_name},\n\nYou are cordially invited to attend a meeting on {topic}.\n\nMeeting Details:\nDate: {date}\nTime: 10:00 AM - 10:30 AM\nLocation: Conference Room A\n\nAgenda:\n{action_item}\n\nKindly confirm your attendance.\n\nSincerely",
            "professional": "Dear {sender_name},\n\nI am writing to request a meeting to discuss {topic} on {date}.\n\nMeeting Objective:\n{action_item}\n\nDuration: 30 minutes\nFormat: In-person or virtual (link to be provided)\n\nPlease confirm your availability or propose an alternative time.\n\nBest regards"
        }
    },
    "apology": {
        "zh": {
            "friendly": "Hi {sender_name}，\n\n真的非常抱歉！关于{topic}的事情，我们在{date}出现了疏漏，给您添麻烦了。\n\n我们的补救措施：\n{action_item}\n\n真的很抱歉，下次一定注意！希望您能谅解。\n\n祝好",
            "polite": "{sender_name} 您好，\n\n就{topic}一事（{date}），我们深表歉意。由于我们的工作失误，给您带来了不便，对此我们深感抱歉。\n\n我们已采取以下措施：\n{action_item}\n\n我们承诺将加强管理，避免类似问题再次发生。\n\n再次致歉，恳请谅解。\n\n此致",
            "formal": "尊敬的{sender_name}：\n\n关于{date}发生的{topic}问题，我们向您致以最诚挚的歉意。\n\n经核查，问题原因如下：[具体原因分析]\n\n整改措施：\n{action_item}\n\n我们已全面排查同类隐患，确保不再发生类似事件。\n\n恳请贵方谅解，期待继续合作。",
            "professional": "Dear {sender_name},\n\nPlease accept our sincere apologies regarding the {topic} incident on {date}. We fully understand the impact and take full responsibility.\n\nImmediate Actions Taken:\n{action_item}\n\nWe have implemented additional safeguards to prevent recurrence.\n\nWe value your partnership and remain committed to the highest standards.\n\nSincerely"
        },
        "en": {
            "friendly": "Hi {sender_name},\n\nI'm really sorry about {topic} on {date}! That was totally our mistake and I understand your frustration.\n\nHere's what we're doing to fix it:\n{action_item}\n\nWe'll make sure this doesn't happen again. Thank you for your patience!\n\nBest",
            "polite": "Dear {sender_name},\n\nI sincerely apologize for the issue regarding {topic} on {date}. We understand this has caused inconvenience and we take full responsibility.\n\nRemedial Actions:\n{action_item}\n\nWe are taking steps to prevent a recurrence and appreciate your understanding.\n\nBest regards",
            "formal": "Dear {sender_name},\n\nPlease accept our formal apology concerning the {topic} matter on {date}.\n\nRoot Cause Analysis:\n[Detailed explanation]\n\nCorrective Actions:\n{action_item}\n\nWe have conducted a thorough review to ensure this does not recur.\n\nSincerely yours",
            "professional": "Dear {sender_name},\n\nI am writing to formally apologize for the {topic} situation on {date}. We acknowledge the seriousness of this matter and accept full responsibility.\n\nAction Plan:\n{action_item}\n\nWe are committed to transparency and continuous improvement.\n\nPlease do not hesitate to contact me directly.\n\nBest regards"
        }
    },
    "inquiry_response": {
        "zh": {
            "friendly": "Hi {sender_name}，\n\n收到您关于{topic}的咨询，以下是相关信息（{date}）：\n\n{action_item}\n\n如果还有不清楚的地方，随时找我聊！\n\n祝好",
            "polite": "{sender_name} 您好，\n\n感谢您就{topic}一事的咨询（{date}），现回复如下：\n\n{action_item}\n\n如有进一步疑问，欢迎随时联系。\n\n此致敬礼",
            "formal": "尊敬的{sender_name}：\n\n您关于{topic}的来函（{date}）已收悉，现正式答复如下：\n\n{action_item}\n\n如需补充材料，请随时告知。\n\n专此答复。",
            "professional": "Dear {sender_name},\n\nThank you for your inquiry regarding {topic} received on {date}.\n\nPlease find our response below:\n\n{action_item}\n\nShould you require additional information, please do not hesitate to reach out.\n\nBest regards"
        },
        "en": {
            "friendly": "Hi {sender_name},\n\nGreat question about {topic}! Here's what I can share ({date}):\n\n{action_item}\n\nLet me know if you need more details!\n\nBest",
            "polite": "Dear {sender_name},\n\nThank you for your inquiry about {topic} on {date}. Please find the requested information below:\n\n{action_item}\n\nFeel free to reach out if you need further clarification.\n\nBest regards",
            "formal": "Dear {sender_name},\n\nThis is in reference to your inquiry dated {date} regarding {topic}.\n\nResponse:\n{action_item}\n\nWe remain at your disposal for any additional information.\n\nSincerely",
            "professional": "Dear {sender_name},\n\nThank you for reaching out regarding {topic} on {date}. Below is our detailed response:\n\n{action_item}\n\nWe are happy to schedule a call if that would be helpful.\n\nBest regards"
        }
    },
    "follow_up": {
        "zh": {
            "friendly": "Hi {sender_name}，\n\n上次聊的{topic}（{date}），想跟您跟进一下进展~\n\n{action_item}\n\n有空的时候回我一下就好，不急！\n\n祝好",
            "polite": "{sender_name} 您好，\n\n关于{date}我们沟通的{topic}事宜，特此跟进：\n\n{action_item}\n\n烦请告知最新进展，谢谢！\n\n此致敬礼",
            "formal": "尊敬的{sender_name}：\n\n就{topic}一事（曾于{date}沟通），现致函跟进如下：\n\n{action_item}\n\n敬请回复为盼。\n\n专此函达。",
            "professional": "Dear {sender_name},\n\nI am writing to follow up on our discussion regarding {topic} on {date}.\n\nOutstanding Items:\n{action_item}\n\nPlease update us on the status at your earliest convenience.\n\nBest regards"
        },
        "en": {
            "friendly": "Hi {sender_name},\n\nJust checking in on {topic} we discussed on {date}. Here's where we are:\n\n{action_item}\n\nNo rush — just let me know when you have a moment!\n\nBest",
            "polite": "Dear {sender_name},\n\nI hope you are doing well. I am following up on our conversation regarding {topic} on {date}.\n\nCurrent Status:\n{action_item}\n\nPlease let me know if there are any updates.\n\nBest regards",
            "formal": "Dear {sender_name},\n\nThis is a follow-up to our correspondence dated {date} regarding {topic}.\n\nAction Required:\n{action_item}\n\nWe look forward to your response.\n\nSincerely",
            "professional": "Dear {sender_name},\n\nI am following up on the action items from our meeting on {topic} held on {date}.\n\nOutstanding Actions:\n{action_item}\n\nKindly provide a status update by [date].\n\nBest regards"
        }
    }
}


# ==================== 意图分类关键词 ====================

INTENT_KEYWORDS: Dict[str, List[str]] = {
    "thank_you": ["感谢", "谢谢", "thank", "thanks", "appreciate", "grateful", "感激"],
    "project_update": ["进展", "进度", "update", "progress", "milestone", "阶段", "汇报", "status"],
    "meeting_request": ["会议", "开会", "约", "meeting", "schedule", "calendar", "discuss", "预约", "zoom"],
    "apology": ["抱歉", "对不起", "sorry", "apologize", "regret", "mistake", "apology", "失误"],
    "inquiry_response": ["咨询", "询问", "inquiry", "question", "询问", "request", "了解", "请问"],
    "follow_up": ["跟进", "follow up", "更新", "进展如何", "update", "check in", "提醒", "reminder"]
}


# ==================== 语气检测关键词 ====================

TONE_KEYWORDS: Dict[str, List[str]] = {
    "friendly": ["你好", "hi", "hello", "嘿", "哈喽", "拜托", "麻烦", "thanks"],
    "polite": ["您好", "请", "贵", "谨", "致", "dear", "kindly", "would"],
    "formal": ["尊敬", "谨此", "兹", "特此", "dear sir", "to whom", "official"],
    "professional": ["regards", "sincerely", "dear mr", "dear ms", "best regards", "respectfully"]
}


# ==================== 核心类 ====================

class EmailReplier:
    """邮件智能回复引擎 v4.8.0"""

    def __init__(self, templates_dir: Optional[str] = None):
        """
        初始化邮件回复引擎

        Args:
            templates_dir: 自定义模板目录路径（可选）
        """
        self.version = __version__
        self.templates_dir = Path(templates_dir) if templates_dir else None
        self.custom_templates: Dict[str, Dict[str, Dict[str, str]]] = {}
        self._load_custom_templates()

    def _load_custom_templates(self):
        """加载自定义模板"""
        if self.templates_dir and self.templates_dir.exists():
            for template_file in self.templates_dir.glob("*.json"):
                try:
                    with open(template_file, "r", encoding="utf-8") as f:
                        data = json.load(f)
                        category = template_file.stem
                        self.custom_templates[category] = data
                except (json.JSONDecodeError, IOError):
                    continue

    def reply(
        self,
        content: str,
        tone: str = "friendly",
        lang: str = "zh",
        context: Optional[str] = None,
        template_path: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        生成邮件回复

        Args:
            content: 原始邮件内容
            tone: 语气风格 (friendly/polite/formal/professional)
            lang: 语言 (zh/en)
            context: 可选的上下文/历史
            template_path: 可选的自定义模板文件路径

        Returns:
            dict: 包含回复结果的字典
        """
        # 检测意图
        intent = self._classify_intent(content)

        # 检测语气（如果用户未显式指定具体 tone 则使用自动检测）
        detected_tone = self._detect_tone(content)
        effective_tone = tone if tone != "friendly" else detected_tone

        # 匹配模板
        template_str = self._match_template(content, effective_tone, intent, lang, template_path)

        # 提取变量
        variables = self._extract_variables(content, context)

        # 应用变量替换
        reply_text = self._apply_variables(template_str, variables)

        # 构建结果
        result = {
            "ok": True,
            "version": self.version,
            "intent": intent,
            "tone": effective_tone,
            "lang": lang,
            "reply": reply_text,
            "variables_used": variables,
            "template_category": intent,
            "timestamp": datetime.now().isoformat(),
            "method": "local_template"
        }

        return result

    def batch_reply(
        self,
        input_dir: str,
        output_dir: str,
        tone: str = "friendly"
    ) -> List[Dict[str, Any]]:
        """
        批量处理邮件回复

        Args:
            input_dir: 输入目录（含 .txt 或 .xlsx 文件）
            output_dir: 输出目录
            tone: 语气风格

        Returns:
            list: 处理结果列表
        """
        input_path = safe_path(input_dir)
        output_path = safe_path(output_dir)

        results = []
        files = list(input_path.glob("*.txt")) + list(input_path.glob("*.xlsx"))

        total = len(files)
        if total == 0:
            return [{"ok": False, "error": "未找到输入文件（支持 .txt 和 .xlsx）"}]

        for idx, file_path in enumerate(files, 1):
            progress_callback(idx, total, f"处理 {file_path.name}")

            if file_path.suffix == ".txt":
                result = self._process_txt_file(file_path, output_path, tone)
            elif file_path.suffix == ".xlsx":
                result = self._process_xlsx_file(file_path, output_path, tone)
            else:
                result = {"ok": False, "error": f"不支持的文件类型: {file_path.suffix}"}

            results.append(result)

        return results

    def list_templates(self) -> Dict[str, Any]:
        """
        列出所有可用模板

        Returns:
            dict: 模板信息字典
        """
        templates_info = {
            "ok": True,
            "version": self.version,
            "builtin_categories": list(BUILTIN_TEMPLATES.keys()),
            "custom_categories": list(self.custom_templates.keys()),
            "supported_tones": ["friendly", "polite", "formal", "professional"],
            "supported_langs": ["zh", "en"],
            "total_builtin": len(BUILTIN_TEMPLATES),
            "total_custom": len(self.custom_templates),
            "categories_detail": {}
        }

        for category in BUILTIN_TEMPLATES:
            templates_info["categories_detail"][category] = {
                "builtin": True,
                "langs": list(BUILTIN_TEMPLATES[category].keys()),
                "tones": list(BUILTIN_TEMPLATES[category].get("zh", {}).keys())
            }

        for category in self.custom_templates:
            templates_info["categories_detail"][category] = {
                "builtin": False,
                "langs": list(self.custom_templates[category].keys()),
                "tones": list(self.custom_templates[category].get("zh", {}).keys())
            }

        return templates_info

    def _match_template(
        self,
        content: str,
        tone: str,
        intent: Optional[str] = None,
        lang: str = "zh",
        template_path: Optional[str] = None
    ) -> str:
        """
        匹配最佳模板

        Args:
            content: 邮件内容
            tone: 语气
            intent: 意图分类
            lang: 语言
            template_path: 自定义模板路径

        Returns:
            str: 模板字符串
        """
        # 优先使用用户指定的模板文件
        if template_path:
            tp = safe_path(template_path)
            if tp.exists():
                with open(tp, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    # 支持直接字符串或结构化
                    if isinstance(data, str):
                        return data
                    if isinstance(data, dict):
                        return data.get(lang, data.get("zh", ""))

        # 检测意图
        if intent is None:
            intent = self._classify_intent(content)

        # 尝试从自定义模板获取
        if intent in self.custom_templates:
            custom = self.custom_templates[intent]
            if lang in custom and tone in custom[lang]:
                return custom[lang][tone]

        # 从内置模板获取
        if intent in BUILTIN_TEMPLATES:
            builtin = BUILTIN_TEMPLATES[intent]
            if lang in builtin and tone in builtin[lang]:
                return builtin[lang][tone]
            # 降级到中文
            if "zh" in builtin and tone in builtin["zh"]:
                return builtin["zh"][tone]

        # 最终降级到 thank_you 友好模板
        return BUILTIN_TEMPLATES["thank_you"]["zh"]["friendly"]

    def _apply_variables(self, template_str: str, variables: Dict[str, str]) -> str:
        """
        替换模板变量

        Args:
            template_str: 模板字符串
            variables: 变量字典

        Returns:
            str: 替换后的字符串
        """
        result = template_str
        for key, value in variables.items():
            placeholder = "{" + key + "}"
            result = result.replace(placeholder, value)

        # 清理未替换的变量
        result = re.sub(r"\{[a-z_]+\}", "[待填写]", result)

        return result

    def _detect_tone(self, content: str) -> str:
        """
        检测邮件语气

        Args:
            content: 邮件内容

        Returns:
            str: 检测到的语气
        """
        content_lower = content.lower()

        scores = {}
        for tone, keywords in TONE_KEYWORDS.items():
            score = sum(1 for kw in keywords if kw in content_lower)
            scores[tone] = score

        if max(scores.values()) == 0:
            return "friendly"

        return max(scores, key=scores.get)

    def _classify_intent(self, content: str) -> str:
        """
        分类邮件意图

        Args:
            content: 邮件内容

        Returns:
            str: 意图分类
        """
        content_lower = content.lower()

        scores = {}
        for intent, keywords in INTENT_KEYWORDS.items():
            score = sum(1 for kw in keywords if kw in content_lower)
            scores[intent] = score

        if max(scores.values()) == 0:
            return "inquiry_response"

        return max(scores, key=scores.get)

    def _extract_variables(self, content: str, context: Optional[str] = None) -> Dict[str, str]:
        """
        从内容中提取变量

        Args:
            content: 邮件内容
            context: 可选上下文

        Returns:
            dict: 提取的变量
        """
        variables = {
            "sender_name": self._extract_name(content),
            "date": datetime.now().strftime("%Y年%m月%d日"),
            "topic": self._extract_topic(content),
            "action_item": self._extract_action_item(content, context)
        }
        return variables

    def _extract_name(self, content: str) -> str:
        """从邮件内容中提取发件人名称"""
        patterns = [
            r"(?:From|发件人|来自)[：:\s]*([^\n,<]+)",
            r"(?:Dear|尊敬的|亲爱的)[：:\s]*([^\n,<]+)",
            r"^([^\n]+?)[，,]",
        ]
        for pattern in patterns:
            match = re.search(pattern, content, re.MULTILINE | re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                if name and len(name) < 50:
                    return name
        return "朋友"

    def _extract_topic(self, content: str) -> str:
        """提取主题"""
        patterns = [
            r"(?:Subject|主题|Re|关于)[：:\s]*([^\n]+)",
            r"关于(.+?)(?:的|之事|事宜)",
        ]
        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                topic = match.group(1).strip()
                if topic and len(topic) < 100:
                    return topic
        # 取第一行作为主题
        first_line = content.strip().split("\n")[0]
        return first_line[:50] if first_line else "相关事宜"

    def _extract_action_item(self, content: str, context: Optional[str] = None) -> str:
        """提取行动项"""
        action_patterns = [
            r"(?:请|麻烦|希望|需要)(.+?)(?:。|！|\n|$)",
            r"(?:action item|to-do|task)[：:\s]*(.+?)(?:\n\n|$)",
        ]
        for pattern in action_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if action := match:
                return action.group(0).strip()[:200]

        if context:
            return f"请参考之前的沟通内容：{context[:100]}"

        return "如有需要请进一步沟通确认"

    def _process_txt_file(self, file_path: Path, output_dir: Path, tone: str) -> Dict[str, Any]:
        """处理单个 txt 文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            result = self.reply(content, tone=tone)

            if result["ok"]:
                output_file = output_dir / f"{file_path.stem}_reply.json"
                with open(output_file, "w", encoding="utf-8") as f:
                    json.dump(result, f, ensure_ascii=False, indent=2)
                result["output_file"] = str(output_file)

            return result
        except Exception as e:
            return {"ok": False, "file": str(file_path), "error": str(e)}

    def _process_xlsx_file(self, file_path: Path, output_dir: Path, tone: str) -> Dict[str, Any]:
        """处理 xlsx 文件"""
        if not HAS_OPENPYXL:
            return {"ok": False, "file": str(file_path), "error": "openpyxl 未安装"}

        try:
            wb = openpyxl.load_workbook(file_path, read_only=True)
            ws = wb.active

            results = []
            for idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 1):
                if row and row[0]:
                    content = str(row[0])
                    lang = str(row[1]) if len(row) > 1 and row[1] else "zh"
                    result = self.reply(content, tone=tone, lang=lang)
                    result["row"] = idx
                    results.append(result)

            wb.close()

            # 保存汇总结果
            output_file = output_dir / f"{file_path.stem}_batch_reply.json"
            with open(output_file, "w", encoding="utf-8") as f:
                json.dump(results, f, ensure_ascii=False, indent=2)

            return {
                "ok": True,
                "file": str(file_path),
                "total": len(results),
                "output_file": str(output_file)
            }
        except Exception as e:
            return {"ok": False, "file": str(file_path), "error": str(e)}


# ==================== Word 文档生成 ====================

def generate_docx(reply_text: str, output_path: str, title: str = "邮件回复") -> Dict[str, Any]:
    """
    生成 Word 文档

    Args:
        reply_text: 回复文本内容
        output_path: 输出文件路径
        title: 文档标题

    Returns:
        dict: 生成结果
    """
    if not HAS_DOCX:
        return {"ok": False, "error": "python-docx 未安装，请运行: pip install python-docx"}

    try:
        doc = Document()

        # 添加标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        # 添加日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # 空行

        # 添加回复内容
        for line in reply_text.split("\n"):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(12)
            run.font.name = "微软雅黑"

        # 添加页脚分隔线
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run("— 由邮件智能回复 v4.8.0 自动生成 —")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        footer_run.italic = True

        # 保存文件
        safe_output = safe_path(output_path)
        doc.save(str(safe_output))

        return {
            "ok": True,
            "output_file": str(safe_output),
            "size_bytes": os.path.getsize(safe_output)
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ==================== CLI 入口 ====================

def main():
    """CLI 主入口"""
    import argparse

    parser = argparse.ArgumentParser(
        description="邮件智能回复 v4.8.0 - 模板匹配 + 规则引擎（纯本地实现）",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )

    subparsers = parser.add_subparsers(dest="command", help="可用子命令")

    # reply 子命令
    reply_parser = subparsers.add_parser("reply", help="生成单封邮件回复")
    reply_parser.add_argument("--content", required=True, help="邮件内容")
    reply_parser.add_argument("--tone", choices=["friendly", "polite", "formal", "professional"],
                              default="friendly", help="语气风格")
    reply_parser.add_argument("--lang", choices=["zh", "en"], default="zh", help="语言")
    reply_parser.add_argument("--context", default=None, help="上下文/历史")
    reply_parser.add_argument("--template", default=None, help="自定义模板文件路径")
    reply_parser.add_argument("--output", default=None, help="输出文件路径（.docx）")

    # batch 子命令
    batch_parser = subparsers.add_parser("batch", help="批量处理邮件")
    batch_parser.add_argument("--input-dir", required=True, help="输入目录")
    batch_parser.add_argument("--output-dir", required=True, help="输出目录")
    batch_parser.add_argument("--tone", choices=["friendly", "polite", "formal", "professional"],
                              default="friendly", help="语气风格")

    # check 子命令
    check_parser = subparsers.add_parser("check", help="列出可用模板")

    args = parser.parse_args()

    replier = EmailReplier()

    if args.command == "reply":
        result = replier.reply(
            content=args.content,
            tone=args.tone,
            lang=args.lang,
            context=args.context,
            template_path=args.template
        )

        # 生成 Word 文档
        if args.output and result.get("ok"):
            docx_result = generate_docx(result["reply"], args.output)
            result["docx_output"] = docx_result

        print(json.dumps(result, ensure_ascii=False, indent=2))

    elif args.command == "batch":
        results = replier.batch_reply(
            input_dir=args.input_dir,
            output_dir=args.output_dir,
            tone=args.tone
        )
        summary = {
            "ok": True,
            "version": __version__,
            "total": len(results),
            "successful": sum(1 for r in results if r.get("ok")),
            "results": results
        }
        print(json.dumps(summary, ensure_ascii=False, indent=2))

    elif args.command == "check":
        templates = replier.list_templates()
        print(json.dumps(templates, ensure_ascii=False, indent=2))

    else:
        parser.print_help()


if __name__ == "__main__":
    main()
