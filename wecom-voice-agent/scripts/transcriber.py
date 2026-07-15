#!/usr/bin/env python3
"""
企业微信语音助手 — 通话录音文字转写全文 v2.0

功能：
1. 通话结束后，将 ASR 完整文字流输出为 .txt（标准库）
2. 支持输出为 .docx（需安装 python-docx，可选）
3. 按说话人分离、时间戳标注
4. 保存到本地目录

特性：
- .txt 输出纯 Python 标准库，零依赖
- .docx 输出可选支持（未安装 python-docx 时降级为 .txt）
- 支持按角色（user/agent）分离输出
- 支持添加时间戳

更新日志：
| 版本 | 日期 | 更新内容 |
|------|------|----------|
| v2.0 | 2026-07-15 | 初始发布：TXT/DOCX 输出、角色分离、时间戳 |
| v2.0.1 | 2026-07-15 | 路径遍历防护、中文化日志、异常处理增强 |

联系信息：njskills@agent.qq.com

使用方法：
    python scripts/transcriber.py                        # 自测
    python scripts/transcriber.py --input turns.json     # 从 JSON 文件读取
    python scripts/transcriber.py --output minutes/      # 指定输出目录
"""

import json
import logging
import os
import re
import sys
from datetime import datetime
from typing import Dict, List, Optional

# ==========================================
# 配置
# ==========================================

OUTPUT_DIR = os.path.join(os.path.expanduser("~"), ".wecom_voice", "transcripts")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)
logger = logging.getLogger("Transcriber")

# 尝试导入 python-docx（可选）
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.info("python-docx 未安装，将仅输出 .txt 格式。安装命令: pip install python-docx")


def _sanitize_filename(filename: str) -> str:
    """
    清理文件名，防止路径遍历攻击。
    仅允许字母、数字、下划线、连字符和点号。
    """
    if not filename or not isinstance(filename, str):
        raise ValueError("文件名不能为空且必须是字符串")
    # 仅允许安全字符（去掉路径分隔符）
    if not re.match(r'^[a-zA-Z0-9_\-\.]+$', filename):
        raise ValueError(f"文件名包含非法字符，仅允许字母、数字、下划线、连字符和点号: {filename}")
    # 不允许以 . 开头（隐藏文件）
    if filename.startswith('.'):
        raise ValueError(f"文件名不能以点号开头: {filename}")
    return filename


# ==========================================
# 转写器
# ==========================================

class TranscriptWriter:
    """
    通话全文转写器

    使用方式：
        writer = TranscriptWriter()
        turns = [{"role": "user", "content": "...", "time": "..."}, ...]
        writer.write_txt(turns, "call_001.txt")
        writer.write_docx(turns, "call_001.docx")
    """

    def __init__(self, output_dir: str = OUTPUT_DIR):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def write_txt(self, turns: List[Dict], filename: str = None,
                  call_id: str = "unknown", include_timestamp: bool = True,
                  include_speaker: bool = True) -> str:
        """
        输出为 .txt 文件

        Args:
            turns: 对话轮次列表，格式 [{"role": "user"/"agent", "content": "...", "time": "..."}]
            filename: 输出文件名（None 则自动命名为 call_id.txt）
            call_id: 通话ID
            include_timestamp: 是否包含时间戳
            include_speaker: 是否标注说话人

        Returns:
            str: 输出文件路径
        """
        if not filename:
            filename = f"{call_id}.txt"

        # 安全性：校验文件名，防止路径遍历
        safe_filename = _sanitize_filename(filename)
        filepath = os.path.join(self.output_dir, safe_filename)

        # 二次确保解析后的路径仍在 output_dir 下
        if not os.path.abspath(filepath).startswith(os.path.abspath(self.output_dir)):
            logger.error(f"文件路径超出允许目录: {filepath}")
            raise ValueError("文件路径不合法：试图写入输出目录之外的位置")

        lines = []
        lines.append("=" * 60)
        lines.append("通话全文转写记录")
        lines.append(f"通话ID: {call_id}")
        lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("=" * 60)
        lines.append("")

        for turn in turns:
            role = turn.get("role", "unknown")
            content = turn.get("content", "")
            timestamp = turn.get("time", "")
            confidence = turn.get("confidence", 0)

            if include_speaker:
                if role == "user":
                    speaker = "👤 用户"
                elif role == "agent":
                    speaker = "🤖 助手"
                else:
                    speaker = f"❓ {role}"
            else:
                speaker = ""

            if include_timestamp and timestamp:
                line = f"[{timestamp}] {speaker}: {content}"
            else:
                line = f"{speaker}: {content}" if speaker else content

            lines.append(line)

        lines.append("")
        lines.append("=" * 60)
        lines.append(f"共 {len(turns)} 条记录")
        lines.append("=" * 60)

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))

        logger.info(f"TXT 已保存: {filepath}")
        return filepath

    def write_docx(self, turns: List[Dict], filename: str = None,
                   call_id: str = "unknown", title: str = "通话全文转写记录") -> Optional[str]:
        """
        输出为 .docx 文件（需要 python-docx）

        Args:
            turns: 对话轮次列表
            filename: 输出文件名
            call_id: 通话ID
            title: 文档标题

        Returns:
            str: 输出文件路径，None 表示未安装 python-docx
        """
        if not HAS_DOCX:
            logger.warning("python-docx 未安装，无法输出 .docx。已降级为 .txt")
            txt_filename = (filename or f"{call_id}.docx").replace(".docx", ".txt")
            return self.write_txt(turns, filename=txt_filename, call_id=call_id)

        if not filename:
            filename = f"{call_id}.docx"

        # 安全性：校验文件名
        safe_filename = _sanitize_filename(filename)
        filepath = os.path.join(self.output_dir, safe_filename)

        # 二次确保路径安全
        if not os.path.abspath(filepath).startswith(os.path.abspath(self.output_dir)):
            logger.error(f"文件路径超出允许目录: {filepath}")
            raise ValueError("文件路径不合法：试图写入输出目录之外的位置")

        doc = Document()

        # 标题
        title_para = doc.add_heading(title, level=0)

        # 元数据
        meta_para = doc.add_paragraph()
        meta_para.add_run("通话ID: ").bold = True
        meta_para.add_run(f"{call_id}\n")
        meta_para.add_run("生成时间: ").bold = True
        meta_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta_para.add_run("记录条数: ").bold = True
        meta_para.add_run(f"{len(turns)}\n")

        doc.add_heading("对话记录", level=1)

        for turn in turns:
            role = turn.get("role", "unknown")
            content = turn.get("content", "")
            timestamp = turn.get("time", "")

            para = doc.add_paragraph()

            if timestamp:
                run = para.add_run(f"[{timestamp}] ")
                run.font.size = Pt(9)
                run.font.color.rgb = None  # 灰色

            if role == "user":
                speaker_run = para.add_run("👤 用户: ")
                speaker_run.bold = True
            elif role == "agent":
                speaker_run = para.add_run("🤖 助手: ")
                speaker_run.bold = True
            else:
                para.add_run(f"❓ {role}: ")

            para.add_run(content)

        # 页脚
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(f"— 共 {len(turns)} 条记录 —")
        footer_run.font.size = Pt(9)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(filepath)
        logger.info(f"DOCX 已保存: {filepath}")
        return filepath

    def write_both(self, turns: List[Dict], call_id: str = "unknown",
                   base_filename: str = None) -> Dict[str, str]:
        """
        同时输出 .txt 和 .docx

        Args:
            turns: 对话轮次列表
            call_id: 通话ID
            base_filename: 基本文件名（不含扩展名）

        Returns:
            dict: {"txt": 路径, "docx": 路径}
        """
        if not base_filename:
            base_filename = call_id

        result = {}

        txt_path = self.write_txt(turns, filename=f"{base_filename}.txt", call_id=call_id)
        result["txt"] = txt_path

        docx_path = self.write_docx(turns, filename=f"{base_filename}.docx", call_id=call_id)
        result["docx"] = docx_path

        return result

    def list_transcripts(self) -> List[str]:
        """列出所有转写文件"""
        try:
            files = os.listdir(self.output_dir)
            return sorted([f for f in files if f.endswith(('.txt', '.docx'))])
        except Exception as e:
            logger.warning(f"列出转写文件失败: {e}")
            return []

    def read_transcript(self, filename: str) -> Optional[str]:
        """
        读取转写文件内容（带路径遍历防护）

        Args:
            filename: 文件名（不含路径）

        Returns:
            str: 文件内容，None 表示文件不存在
        """
        try:
            safe_filename = _sanitize_filename(filename)
        except ValueError as e:
            logger.warning(f"读取转写文件失败: {e}")
            return None

        filepath = os.path.join(self.output_dir, safe_filename)

        # 二次确保路径安全（防御符号链接攻击）
        if not os.path.abspath(filepath).startswith(os.path.abspath(self.output_dir)):
            logger.warning(f"路径超出允许目录: {filepath}")
            return None

        if not os.path.exists(filepath):
            return None

        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.warning(f"读取文件失败: {e}")
            return None


# ==========================================
# 格式转换工具
# ==========================================

class TurnFormatter:
    """轮次数据格式化工具"""

    @staticmethod
    def from_ivr_format(turns: List[Dict]) -> List[Dict]:
        """
        从 IVR 纪要格式转换（兼容 state_machine 输出）

        输入格式: [{"role": "user", "content": "...", "time": "...", "confidence": 0.95}]
        输出格式: 同上（兼容）
        """
        return turns

    @staticmethod
    def from_raw_text(raw_text: str, split_char: str = "\n") -> List[Dict]:
        """
        从原始文本转换（每行一条记录）

        Args:
            raw_text: 原始转写文本
            split_char: 分隔符（默认换行）
        """
        turns = []
        for i, line in enumerate(raw_text.split(split_char)):
            line = line.strip()
            if not line:
                continue
            turns.append({
                "role": "user" if i % 2 == 0 else "agent",
                "content": line,
                "time": "",
                "confidence": 0,
            })
        return turns

    @staticmethod
    def from_json(filepath: str) -> List[Dict]:
        """从 JSON 文件读取轮次数据"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            if isinstance(data, list):
                return data
            if isinstance(data, dict) and "turns" in data:
                return data["turns"]
            return []
        except Exception as e:
            logger.error(f"JSON 读取失败: {e}")
            return []


# ==========================================
# 自测
# ==========================================

def run_self_test():
    """运行转写器自测"""
    print("=" * 60)
    print("通话全文转写 — 自测模式")
    print("=" * 60)

    # 模拟通话数据
    turns = [
        {"role": "agent", "content": "您好，我是语音助手，请问有什么可以帮您？", "time": "2026-07-15T10:00:00"},
        {"role": "user", "content": "我想查一下明天的天气怎么样", "time": "2026-07-15T10:00:05", "confidence": 0.92},
        {"role": "agent", "content": "明天北京晴，温度25到30度", "time": "2026-07-15T10:00:10"},
        {"role": "user", "content": "好的我知道了谢谢", "time": "2026-07-15T10:00:15", "confidence": 0.95},
        {"role": "agent", "content": "不客气，请问还有其他需要帮助的吗？", "time": "2026-07-15T10:00:20"},
        {"role": "user", "content": "没有了再见", "time": "2026-07-15T10:00:25", "confidence": 0.88},
        {"role": "agent", "content": "再见，祝您生活愉快！", "time": "2026-07-15T10:00:30"},
    ]

    # 测试 1: TXT 输出
    print("\n[测试 1] TXT 输出")
    writer = TranscriptWriter()
    txt_path = writer.write_txt(turns, call_id="call_test_001")
    assert os.path.exists(txt_path)

    content = writer.read_transcript(os.path.basename(txt_path))
    assert "👤 用户" in content and "🤖 助手" in content
    assert "我想查一下明天的天气怎么样" in content
    print(f"  TXT 路径: {txt_path}")
    print("✅ TXT 输出通过")

    # 测试 2: DOCX 输出
    print("\n[测试 2] DOCX 输出")
    docx_path = writer.write_docx(turns, call_id="call_test_001")
    if HAS_DOCX:
        assert docx_path is not None
        assert os.path.exists(docx_path)
        print(f"  DOCX 路径: {docx_path}")
        print("✅ DOCX 输出通过")
    else:
        # 降级为 TXT
        assert docx_path is not None
        assert docx_path.endswith(".txt")
        print("  python-docx 未安装，已降级为 .txt")
        print("✅ 降级输出通过")

    # 测试 3: 同时输出两种格式
    print("\n[测试 3] 同时输出 TXT + DOCX")
    result = writer.write_both(turns, call_id="call_test_002")
    assert os.path.exists(result["txt"])
    assert os.path.exists(result["docx"])
    print(f"  TXT: {result['txt']}")
    print(f"  DOCX: {result['docx']}")
    print("✅ 同时输出通过")

    # 测试 4: 列出文件
    print("\n[测试 4] 列出转写文件")
    files = writer.list_transcripts()
    assert len(files) >= 2
    print(f"  已有文件: {files}")
    print("✅ 列出文件通过")

    # 测试 5: 从原始文本转换
    print("\n[测试 5] 从原始文本转换")
    raw_text = "你好\n你好请问有什么需要帮忙的\n我想订机票\n请问你要订哪天的\n明天的\n好的"
    raw_turns = TurnFormatter.from_raw_text(raw_text)
    assert len(raw_turns) == 6
    assert raw_turns[0]["role"] == "user"
    assert raw_turns[1]["role"] == "agent"
    print(f"  转换得到 {len(raw_turns)} 条记录")
    print("✅ 原始文本转换通过")

    # 测试 6: 无时间戳模式
    print("\n[测试 6] 无时间戳纯文本输出")
    simple_path = writer.write_txt(turns, call_id="call_test_003",
                                   include_timestamp=False, include_speaker=False)
    simple_content = writer.read_transcript(os.path.basename(simple_path))
    assert "👤" not in simple_content  # 无说话人标记
    print("✅ 无时间戳纯文本通过")

    # 测试 7: 路径遍历防护
    print("\n[测试 7] 路径遍历防护")
    bad_content = writer.read_transcript("../../etc/passwd")
    assert bad_content is None
    try:
        writer.write_txt(turns, filename="../../../tmp/evil.txt")
        print("❌ 应拒绝恶意文件名")
    except ValueError:
        print("✅ 已拦截恶意文件名写入")

    print(f"\n{'='*60}")
    print("所有自测通过 ✓")
    print('='*60)


if __name__ == "__main__":
    run_self_test()
